import json
from fpdf import FPDF
import docx
from json2xml import json2xml
from json2xml.utils import readfromjson
import Basic_System_Details
import DMI_File_Info
import CPU_Info
import Disk_Info
import GPU_Info
import Memory_Info
import Network_Info
import Peripherals_Info
import Software_Info
import System_Uptime
import Users_Info

json_file_path = 'report.json'


def append_to_json(keys: list) -> None:
    audit_dct = {
        "Общие данные о системе": Basic_System_Details.get_Common_Basic_System_Details(),
        "Данные из dmi/id":DMI_File_Info.get_Unix_DMI_id_Info(),
        "Системное время": System_Uptime.get_Unix_System_Uptime_Info(),
        "Пользователи": Users_Info.get_Common_Users_Info(),
        "Ядра ОС": CPU_Info.get_Common_Cores_Info(),
        "Процессоры": CPU_Info.get_Unix_CPU_Info(),
        "HDD/SSD устройства": Disk_Info.get_Common_Disk_Info(),
        "CD/DVD устройства": Disk_Info.get_CD_DVD_info(),
        "USB устройства": Peripherals_Info.get_Unix_USB_Devices(),
        "Данные о GPU": GPU_Info.get_Unix_CGPU_Info(),
        "Данные о RAM памяти": Memory_Info.get_Common_Memory_Info(),
        "Сетевые настройки": Network_Info.get_Common_Network_Info(),
        "Приложения": Software_Info.get_xdg_open_applications(),
        "Исполняемые файлы PATH": Software_Info.get_executables_in_PATH(),
    }
    try:
        report = {k: audit_dct[k] for k in keys}
        # Запись отчёта в JSON файл
        with open(json_file_path, 'w') as file:
            json.dump(report, file, ensure_ascii=False, indent=2)
        print(f"Данные успешно добавлены в файл: {json_file_path}")
    except Exception as e:
        print(f"Произошла ошибка при добавлении данных в файл: {e}")


PROGRAM_NAME = "TotalScan"


# Создание PDF-отчёта
def create_pdf(jsonfile):
    try:
        with open(jsonfile, "r", encoding="utf-8") as file:
            data = json.load(file)
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font('DejaVu', '', 'font/DejaVuSans.ttf', uni=True)
            pdf.set_font("DejaVu", size=20)
            pdf.cell(200, 10, txt=PROGRAM_NAME, ln=1, align="C")
            pdf.set_font("DejaVu", size=16)
            pdf.cell(200, 10, txt="Отчёт аудита безопасности", ln=1, align="C")

            pdf.set_font("DejaVu", size=10)
            for element in data:
                pdf.cell(200, 10, ln=1, txt=element + ": ")
                for part in data[element]:
                    if type(data[element][part]) != dict:
                        pdf.cell(200, 10, ln=1, txt="       " + str(part) + ": " + str(data[element][part]))
                    else:
                        pdf.cell(200, 10, ln=1, txt="       " + str(part) + ": ")
                        for subpart in data[element][part]:
                            if type(data[element][part][subpart]) != dict:
                                pdf.cell(200, 10, ln=1,
                                         txt="               " + str(subpart) + ": " + str(
                                             data[element][part][subpart]))
                            else:
                                pdf.cell(200, 10, ln=1, txt="               " + str(subpart) + ": ")
                                for subsubpart in data[element][part][subpart]:
                                    pdf.cell(200, 10, ln=1,
                                             txt="                       " + str(subsubpart) + ": " + str(
                                                 data[element][part][subpart][subsubpart]))

            pdf.output("output.pdf")
            print(f"PDF-отчёт создан")
    except Exception as e:
        print(f"Произошла ошибка при добавлении данных в файл: {e}")


# Создание HTML-отчёта
def create_html(jsonfile):
    try:
        head = "<html><head><title>" + PROGRAM_NAME + "</title><style>table{border: 1px solid grey;}td{border: 1px solid grey;}</style></head>"
        body = "<body>"
        body += "<h1>TotalScan</h1>"
        body += "<h2>Отчёт об аудите безопасности</h2>"
        with open(jsonfile, "r", encoding="utf-8") as file:
            data = json.load(file)

            for element in data:
                body += "<table><tr><th>" + element + "</th></tr>"
                for part in data[element]:
                    if type(data[element][part]) != dict:
                        body += "<tr><td>" + str(part) + "</td><td>" + str(data[element][part]) + "</td></tr>"
                    else:
                        body += "<tr><td>" + str(part) + "</td><tr>"
                        for subpart in data[element][part]:
                            if type(data[element][part][subpart]) != dict:
                                body += "<td>" + str(subpart) + "</td><td>" + str(
                                    data[element][part][subpart]) + "</td></tr>"
                            else:
                                body += "<td>" + str(subpart) + "</td>"
                                for subsubpart in data[element][part][subpart]:
                                    body += "<td>" + str(subsubpart) + "</td><td>" + str(
                                        data[element][part][subpart][subsubpart]) + "</td>"
                    body += "</tr>"
                body += "</table>"
                body += "<br>"

        html = head + body
        html_output = open("output.html", "w")
        html_output.write(html)
        html_output.close()
        print(f"HTML-отчёт создан")
    except Exception as e:
        print(f"Произошла ошибка при добавлении данных в файл: {e}")


# Создание DOCX-отчёта
def create_docx(jsonfile):
    try:
        doc = docx.Document()
        doc.add_heading(PROGRAM_NAME, 0)

        with open(jsonfile, "r", encoding="utf-8") as file:
            data = json.load(file)

            for element in data:
                doc.add_paragraph(element + ": ")
                for part in data[element]:
                    if type(data[element][part]) != dict:
                        doc.add_paragraph("       " + str(part) + ": " + str(data[element][part]))
                    else:
                        doc.add_paragraph("       " + str(part) + ": ")
                        for subpart in data[element][part]:
                            if type(data[element][part][subpart]) != dict:
                                doc.add_paragraph(
                                    "               " + str(subpart) + ": " + str(data[element][part][subpart]))
                            else:
                                doc.add_paragraph("               " + str(subpart) + ": ")
                                for subsubpart in data[element][part][subpart]:
                                    doc.add_paragraph("                       " + str(subsubpart) + ": " + str(
                                        data[element][part][subpart][subsubpart]))

        doc.save("output.docx")
        print(f"DOCX-отчёт создан")
    except Exception as e:
        print(f"Произошла ошибка при добавлении данных в файл: {e}")


# Создание XML-отчёта
def create_xml(jsonfile):
    try:
        data = readfromjson(jsonfile)
        xml_output = open("output.xml", "w", encoding="utf-8")
        xml_output.write(json2xml.Json2xml(data).to_xml())
        xml_output.close()
        print(f"XML-отчёт создан")
    except Exception as e:
        print(f"Произошла ошибка при добавлении данных в файл: {e}")


def create_reports(keys: list):
    # print(f'keys = {keys}')
    append_to_json(keys)
    create_pdf(json_file_path)
    create_html(json_file_path)
    create_docx(json_file_path)
    create_xml(json_file_path)


if __name__ == '__main__':
    create_reports()
